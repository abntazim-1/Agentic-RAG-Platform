"""
FastAPI + Gradio application — full port from mini_rag.py.

Endpoints (matching mini_rag.py):
  POST /ingest              — ingest text or files
  POST /query               — standard JSON response with timings
  POST /query/stream        — Server-Sent Events token streaming
  POST /evaluate            — batch embedding-based evaluation
  DELETE /source/{name}     — remove a source and its vectors
  GET  /health              — detailed health check

Gradio UI (matching mini_rag.py):
  Tab 1: Chat — streaming, per-session state, timings footer, source citations
  Tab 2: Add Knowledge — multi-file upload, PDF/DOCX/TXT/MD/JSON/CSV support

RAGAS langchain compat patch applied at top (matching mini_rag.py).
"""

# ─── RAGAS PATCH FOR LANGCHAIN 0.4+ COMPATIBILITY ────────────────────────────
# Matches mini_rag.py's monkey-patch exactly
import sys
import types
try:
    import langchain_community.chat_models
    if not hasattr(langchain_community.chat_models, "vertexai"):
        dummy = types.ModuleType("langchain_community.chat_models.vertexai")
        dummy.ChatVertexAI = type("ChatVertexAI", (object,), {})
        sys.modules["langchain_community.chat_models.vertexai"] = dummy
        langchain_community.chat_models.vertexai = dummy
except Exception:
    pass
# ─────────────────────────────────────────────────────────────────────────────

import asyncio
import concurrent.futures as _cf
import json
import os
import time
import uuid

from fastapi import FastAPI, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from dotenv import load_dotenv

load_dotenv()

import sys
import os
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from config import get_settings
from models import QueryRequest, QueryResponse
from ingestion.loaders import load_directory, load_file
from ingestion.chunker import StructureAwareChunker
from ingestion.metadata import MetadataEnricher
from retrieval.embedder import Embedder
from retrieval.vector_store import VectorStore
from retrieval.bm25_store import BM25Store
from retrieval.hybrid_retriever import HybridRetriever
from retrieval.query_rewriter import QueryRewriter
from reranking.reranker import Reranker, ContextCompressor
from memory.conversation import ConversationMemory
from memory.three_tier_memory import ThreeTierMemoryEngine
from gateway.llm_gateway import ResilientLLMGateway
from agent.orchestrator import ProductionAgenticOrchestrator
from evaluation.evaluator import EmbeddingEvaluator
from agent.graph import build_simple_graph
from db.database import init_db, get_db
from api.db_router import router as db_router
from sqlalchemy.orm import Session
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse

# Initialize Database
init_db()

settings = get_settings()

# ─── Singletons (created once, shared across requests) ───────────────────────
embedder       = Embedder()
vector_store   = VectorStore(embedder)
bm25_store     = BM25Store()
query_rewriter = QueryRewriter()
retriever      = HybridRetriever(vector_store, bm25_store, embedder, query_rewriter)
reranker       = Reranker()
compressor     = ContextCompressor()
legacy_memory  = ConversationMemory()
three_tier_memory = ThreeTierMemoryEngine()
llm_gateway    = ResilientLLMGateway()
orchestrator   = ProductionAgenticOrchestrator(retriever, reranker, compressor, three_tier_memory, llm_gateway)
evaluator      = EmbeddingEvaluator(embedder=embedder)

# Rebuild BM25 from chunks restored by VectorStore on startup
bm25_store.build(vector_store.all_chunks)

# Build LangGraph + helpers
rag_graph, run, generate, generate_stream, cached_retrieve_and_rerank = (
    build_simple_graph(retriever, reranker, legacy_memory)
)

# ─── Request models ───────────────────────────────────────────────────────────

class IngestReq(BaseModel):
    text: str
    source: str = "manual"

class QueryReq(BaseModel):
    query: str
    session_id: str = "default"

class EvalReq(BaseModel):
    questions: list[str]
    ground_truths: list[str] = []

# ─── FastAPI app ──────────────────────────────────────────────────────────────
from fastapi.middleware.cors import CORSMiddleware

app = FastAPI(title="Production Agentic RAG Platform")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Include Real-Time Database Inspector Router
app.include_router(db_router)

# Serve Dashboard UI Static Route & File Response
dashboard_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "dashboard")
dashboard_file = os.path.join(dashboard_dir, "index.html")

@app.get("/dashboard")
@app.get("/dashboard/")
def serve_dashboard():
    if os.path.exists(dashboard_file):
        return FileResponse(dashboard_file)
    raise HTTPException(status_code=404, detail="Dashboard index.html file not found")

if os.path.exists(dashboard_dir):
    app.mount("/static_dashboard", StaticFiles(directory=dashboard_dir, html=True), name="static_dashboard")



# ─── /ingest ─────────────────────────────────────────────────────────────────
@app.post("/ingest")
def ingest(req: IngestReq):
    """
    Ingest raw text. Matches mini_rag.py's /ingest endpoint exactly:
      - 409 on duplicate source
      - Chunks + metadata enrichment + embedding + BM25 rebuild
    """
    if req.source in vector_store.ingested_sources:
        raise HTTPException(
            status_code=409,
            detail=f"Source '{req.source}' is already indexed. Delete it first or use a unique source name.",
        )
    try:
        chunker  = StructureAwareChunker()
        enricher = MetadataEnricher()

        # Build a Document-like object for the chunker
        from models import Document
        doc    = Document(content=req.text, source=req.source)
        chunks = chunker.chunk(doc)
        chunks = enricher.enrich_batch(chunks)

        vector_store.add(chunks)
        bm25_store.build(vector_store.all_chunks)
        cached_retrieve_and_rerank.cache_clear()

        return {
            "status":         "ok",
            "source":         req.source,
            "chunks_created": len(chunks),
            "total_chunks":   vector_store.total_chunks,
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ─── /query ──────────────────────────────────────────────────────────────────
@app.post("/query")
def query_endpoint(req: QueryReq):
    """Standard blocking query — returns full JSON response with timings, 3-tier memory save & telemetry."""
    try:
        res = orchestrator.run(req.query, session_id=req.session_id)
        return res.to_dict()
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))



# ─── /query/stream ───────────────────────────────────────────────────────────
@app.post("/query/stream")
async def query_stream(req: QueryReq):
    """
    True token streaming over SSE using the fully Agentic RAG orchestrator loop.
    Runs the agent execution (guardrails, retrieval, self-RAG check, tool fallback, generation, memory)
    in a thread pool and bridges events to an asyncio.Queue to avoid blocking.
    """
    loop = asyncio.get_running_loop()
    q: asyncio.Queue = asyncio.Queue(maxsize=64)
    _DONE = object()

    def _run_orchestrator_stream():
        try:
            for event in orchestrator.run_stream(req.query, session_id=req.session_id):
                loop.call_soon_threadsafe(q.put_nowait, event)
        except Exception as e:
            loop.call_soon_threadsafe(q.put_nowait, {"type": "error", "error": str(e)})
        finally:
            loop.call_soon_threadsafe(q.put_nowait, _DONE)

    async def gen():
        loop.run_in_executor(None, _run_orchestrator_stream)
        
        sources = []
        rewritten_query = req.query
        
        while True:
            item = await q.get()
            if item is _DONE:
                break
                
            if isinstance(item, dict):
                if item.get("type") == "token":
                    yield f"data: {json.dumps({'token': item['token']})}\n\n"
                elif item.get("type") == "status":
                    yield f"data: {json.dumps({'status': item['msg'], 'step': item['step']})}\n\n"
                elif item.get("type") == "error":
                    yield f"data: {json.dumps({'error': item.get('content') or item.get('error')})}\n\n"
                elif item.get("type") == "done":
                    sources = item.get("sources", [])
                    rewritten_query = item.get("rewritten_query", req.query)
            
        yield f"data: {json.dumps({'sources': sources, 'done': True, 'rewritten_query': rewritten_query})}\n\n"

    return StreamingResponse(gen(), media_type="text/event-stream")


# ─── /evaluate ───────────────────────────────────────────────────────────────
@app.post("/evaluate")
def evaluate_endpoint(req: EvalReq):
    """Batch embedding-based evaluation. Matches mini_rag.py's /evaluate."""
    try:
        return evaluator.evaluate_rag(
            run_fn=lambda q: run(q, session_id=f"eval_{uuid.uuid4()}"),
            questions=req.questions,
            ground_truths=req.ground_truths or None,
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ─── DELETE /source/{name} ───────────────────────────────────────────────────
@app.delete("/source/{source_name}")
def delete_source(source_name: str):
    """Remove all vectors + in-memory chunks for a source. Matches mini_rag.py."""
    removed = vector_store.delete_source(source_name)
    bm25_store.build(vector_store.all_chunks)
    cached_retrieve_and_rerank.cache_clear()
    return {"status": "ok", "source": source_name, "chunks_removed": removed}


# ─── /ingest/file ─────────────────────────────────────────────────────────────
@app.post("/ingest/path")
def ingest_path(path: str, save_index: bool = True):
    """Ingest from a filesystem path (file or directory)."""
    if not os.path.exists(path):
        raise HTTPException(status_code=400, detail=f"Path not found: {path}")
    if os.path.isdir(path):
        docs = load_directory(path)
    else:
        docs = [load_file(path)]

    chunker  = StructureAwareChunker()
    enricher = MetadataEnricher()
    all_chunks = []
    for doc in docs:
        if doc.source in vector_store.ingested_sources:
            continue
        chunks = chunker.chunk(doc)
        chunks = enricher.enrich_batch(chunks)
        all_chunks.extend(chunks)

    if all_chunks:
        vector_store.add(all_chunks)
        bm25_store.build(vector_store.all_chunks)
        cached_retrieve_and_rerank.cache_clear()

    return {"status": "ok", "chunks_created": len(all_chunks), "total": vector_store.total_chunks}


# ─── /health ─────────────────────────────────────────────────────────────────
@app.get("/health")
def health():
    """Detailed health check — matches mini_rag.py's /health."""
    info = vector_store.get_collection_info()
    return {
        "status":           "ok",
        "chunks_in_memory": vector_store.total_chunks,
        "vectors_in_qdrant": info.points_count,
        "active_sessions":  legacy_memory.active_sessions,
        "indexed_sources":  list(vector_store.ingested_sources),
    }


# ─── API Endpoints for Eval Dashboard ────────────────────────────────────────

from db.models import QueryMetric, EvalRun
from fastapi import Depends

@app.get("/api/eval/metrics")
def get_metrics(limit: int = 50, db: Session = Depends(get_db)):
    """Fetch recent query metrics with RAGAS scores."""
    metrics = db.query(QueryMetric).order_by(QueryMetric.timestamp.desc()).limit(limit).all()
    return metrics

@app.get("/api/eval/runs")
def get_eval_runs(limit: int = 10, db: Session = Depends(get_db)):
    """Fetch benchmark runs."""
    runs = db.query(EvalRun).order_by(EvalRun.timestamp.desc()).limit(limit).all()
    return runs


# ─── GRADIO UI ───────────────────────────────────────────────────────────────
import gradio as gr

# ─── File text extractor — matches mini_rag.py's extract_text_from_file() ────
def extract_text_from_file(file_path: str) -> tuple[str, str]:
    """Extract text content and derive a source name from an uploaded file."""
    filename    = os.path.basename(file_path)
    source_name = filename
    ext         = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    try:
        if ext in ("txt", "md", "rst", "log", "csv"):
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read(), source_name

        elif ext == "pdf":
            try:
                import pypdf
                reader = pypdf.PdfReader(file_path)
                text   = "\n".join(page.extract_text() or "" for page in reader.pages)
                return text, source_name
            except ImportError:
                return "", "__error__:Install pypdf: pip install pypdf"

        elif ext == "docx":
            try:
                import docx
                doc  = docx.Document(file_path)
                text = "\n".join(p.text for p in doc.paragraphs)
                return text, source_name
            except ImportError:
                return "", "__error__:Install python-docx: pip install python-docx"

        elif ext == "json":
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                import json as _json
                data = _json.load(f)
                return _json.dumps(data, indent=2), source_name

        else:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read(), source_name

    except Exception as e:
        return "", f"__error__:{e}"


def ingest_file(files) -> str:
    """Handle one or more uploaded files. Matches mini_rag.py's ingest_file()."""
    if not files:
        return "⚠️ Please upload at least one file."

    results = []
    for file in (files if isinstance(files, list) else [files]):
        file_path   = file.name if hasattr(file, "name") else str(file)
        text, source_name = extract_text_from_file(file_path)

        if source_name.startswith("__error__:"):
            results.append(f"Error: {os.path.basename(file_path)}: {source_name[10:]}")
            continue

        if not text.strip():
            results.append(f"⚠️ {source_name}: No text could be extracted.")
            continue

        if source_name in vector_store.ingested_sources:
            results.append(f"⚠️ '{source_name}' is already indexed. Delete it first or rename the file.")
            continue

        try:
            from models import Document
            doc    = Document(content=text, source=source_name)
            chunker  = StructureAwareChunker()
            enricher = MetadataEnricher()
            chunks = chunker.chunk(doc)
            chunks = enricher.enrich_batch(chunks)
            vector_store.add(chunks)
            results.append(f"✅ Indexed '{source_name}' → {len(chunks)} chunks (total: {vector_store.total_chunks})")
        except Exception as e:
            results.append(f"❌ Error '{source_name}': {e}")

    # Rebuild BM25 once after all files are processed (PERF 3 fix from mini_rag.py)
    bm25_store.build(vector_store.all_chunks)
    cached_retrieve_and_rerank.cache_clear()
    return "\n".join(results)


def get_indexed_sources_markdown() -> str:
    """Returns formatted markdown table listing all currently uploaded and indexed files."""
    sources = list(vector_store.ingested_sources)
    if not sources:
        return "ℹ️ *No files have been uploaded or indexed yet.*"
    
    md_lines = [
        "| # | Filename / Source Name | Vector Chunks | Status |",
        "|---|---|---|---|"
    ]
    for idx, src in enumerate(sorted(sources)):
        chunk_count = sum(1 for c in vector_store.all_chunks if c.source == src)
        md_lines.append(f"| {idx+1} | `{src}` | **{chunk_count}** chunks | Indexed & Active |")
    
    return "\n".join(md_lines)


def delete_indexed_source_fn(source_name: str) -> tuple[str, str]:
    """Deletes an indexed source and updates the indexed files list."""
    src = source_name.strip()
    if not src:
        return "⚠️ Please enter a valid source name to delete.", get_indexed_sources_markdown()
    
    if src not in vector_store.ingested_sources:
        return f"❌ Source '{src}' not found in index.", get_indexed_sources_markdown()
    
    removed = vector_store.delete_source(src)
    bm25_store.build(vector_store.all_chunks)
    cached_retrieve_and_rerank.cache_clear()
    
    msg = f"✅ Deleted source '{src}' ({removed} chunks removed)."
    return msg, get_indexed_sources_markdown()


def ingest_file_and_refresh(files) -> tuple[str, str]:
    """Ingests uploaded files and updates the indexed files list markdown."""
    status = ingest_file(files)
    updated_sources = get_indexed_sources_markdown()
    return status, updated_sources



# ─── Background RAGAS evaluator pool ─────────────────────────────────────────
_ragas_pool = _cf.ThreadPoolExecutor(max_workers=1, thread_name_prefix="ragas")

def _run_ragas_background(metric_id: str, question: str, answer: str, contexts: list[str]):
    """Fire-and-forget RAGAS. Matches mini_rag.py's _run_ragas_background()."""
    try:
        res = evaluator.evaluate_single(question=question, answer=answer, contexts=contexts)
        
        # Save results to DB
        from db.database import SessionLocal
        from db.models import QueryMetric
        with SessionLocal() as db:
            metric = db.query(QueryMetric).filter(QueryMetric.id == metric_id).first()
            if metric:
                if res.faithfulness is not None:
                    metric.faithfulness = res.faithfulness
                if res.answer_relevancy is not None:
                    metric.answer_relevancy = res.answer_relevancy
                db.commit()
                
        if res.faithfulness is not None:
            parts = []
            if res.faithfulness     is not None: parts.append(f"Faithfulness: {res.faithfulness:.2f}")
            if res.answer_relevancy is not None: parts.append(f"Answer Relevancy: {res.answer_relevancy:.2f}")
            print(f"[RAGAS] {' | '.join(parts)}")
    except Exception as e:
        print(f"[RAGAS] Error: {e}")


def chat_fn(message: str, history: list, session_id: str):
    """
    Streaming chat with 3-tier memory, background RAGAS evaluation, and orchestrator.
    """
    if callable(session_id):
        session_id = session_id()
    if not isinstance(session_id, str):
        session_id = str(session_id)

    try:
        status_steps = []
        answer_parts = []
        sources = []
        rewritten_q = message
        timings = {}
        trace_id = ""

        # Consume the run_stream generator from the agentic orchestrator
        for event in orchestrator.run_stream(message, session_id=session_id):
            if event["type"] == "status":
                status_steps.append(event["msg"])
                current_status = "\n".join(f"*{step}*" for step in status_steps)
                yield f"{current_status}\n\n---"
                
            elif event["type"] == "token":
                answer_parts.append(event["token"])
                current_status = "\n".join(f"*{step}*" for step in status_steps)
                current_answer = "".join(answer_parts)
                yield f"{current_status}\n\n---\n\n{current_answer}"
                
            elif event["type"] == "error":
                yield f"❌ Error: {event.get('content') or event.get('error')}"
                return
                
            elif event["type"] == "done":
                sources = event.get("sources", [])
                timings = event.get("timings", {})
                rewritten_q = event.get("rewritten_query", message)
                trace_id = event.get("trace_id", "")

        answer = "".join(answer_parts).strip()
        if not answer:
            yield "The model returned an empty response. Please try again."
            return

        # ── Footer: timings + rewritten query + sources ──────────────────────
        footer = ""
        timing_parts = []
        if timings.get("rewrite_ms"):  timing_parts.append(f"Rewrite: {timings['rewrite_ms']}ms")
        if timings.get("guardrails_ms"): timing_parts.append(f"Guard: {timings['guardrails_ms']}ms")
        if timings.get("memory_read_ms"): timing_parts.append(f"Mem: {timings['memory_read_ms']}ms")
        if timings.get("retrieval_total_ms"): timing_parts.append(f"Retrieve: {timings['retrieval_total_ms']}ms")
        if timings.get("self_rag_ms"): timing_parts.append(f"Self-RAG: {timings['self_rag_ms']}ms")
        if timings.get("tool_execution_ms"): timing_parts.append(f"Tool: {timings['tool_execution_ms']}ms")
        if timings.get("ttft_ms"): timing_parts.append(f"TTFT: {timings['ttft_ms']}ms")
        if timings.get("generation_ms"): timing_parts.append(f"Gen: {timings['generation_ms']}ms")
        if timings.get("total_latency_ms"): timing_parts.append(f"Total: {timings['total_latency_ms']}ms")
        
        footer += f"\n\n⏱️ **Timings:** `{'  |  '.join(timing_parts)}`"

        if rewritten_q and rewritten_q != message:
            footer += f"\n\n🔍 **Rewritten Query:** `{rewritten_q}`"

        if sources:
            footer += "\n\n**Sources Used:**"
            for s in sources:
                footer += f"\n- [{s['id']}] `{s['source']}` (Confidence: {s['score']})"

        current_status = "\n".join(f"*{step}*" for step in status_steps)
        yield f"{current_status}\n\n---\n\n{answer}{footer}"

        # ── Save to DB for Evaluation Dashboard ───────────────────────────────
        from db.database import SessionLocal
        from db.models import QueryMetric
        
        metric_id = str(uuid.uuid4())
        contexts = [s["content"] for s in sources]
        
        # Convert ms back to seconds for QueryMetric dashboard compatibility
        rewrite_sec = timings.get("rewrite_ms", 0.0) / 1000.0 if timings.get("rewrite_ms") else None
        retrieve_sec = timings.get("retrieval_total_ms", 0.0) / 1000.0 if timings.get("retrieval_total_ms") else None
        ttft_sec = timings.get("ttft_ms", 0.0) / 1000.0 if timings.get("ttft_ms") else None
        gen_sec = timings.get("generation_ms", 0.0) / 1000.0 if timings.get("generation_ms") else None
        total_sec = timings.get("total_latency_ms", 0.0) / 1000.0 if timings.get("total_latency_ms") else None

        try:
            with SessionLocal() as db:
                metric = QueryMetric(
                    id=metric_id,
                    session_id=session_id,
                    query=message,
                    rewritten_query=rewritten_q,
                    answer=answer,
                    rewrite_time=rewrite_sec,
                    retrieve_time=retrieve_sec,
                    ttft=ttft_sec,
                    generation_time=gen_sec,
                    total_time=total_sec,
                    contexts=contexts
                )
                db.add(metric)
                db.commit()
        except Exception as db_e:
            print(f"DB Error: {db_e}")

        # ── Background RAGAS ─────────────────────────────────────────────────
        if contexts and "blocked by" not in answer.lower():
            _ragas_pool.submit(_run_ragas_background, metric_id, message, answer, contexts)

    except Exception as e:
        yield f"❌ Error: {str(e)}"


# ─── Gradio UI blocks — matches mini_rag.py's gr.Blocks() exactly ────────────
with gr.Blocks(title="Production Agentic RAG System") as demo:
    gr.Markdown("# Production-Grade Agentic RAG System")
    gr.Markdown(
        "**Architecture:** Hybrid Search (Dense Qdrant + Sparse BM25) | Cross-Encoder Reranking (TinyBERT) | "
        "3-Tier Memory Store (Working, Semantic, Episodic, Procedural) | Resilient LLM Gateway | OTel Latency Telemetry"
    )
    def get_new_session_id():
        return f"gradio_{uuid.uuid4().hex[:8]}"

    session_state = gr.State(get_new_session_id)


    with gr.Tabs():
        with gr.TabItem("💬 Chat"):
            gr.ChatInterface(
                fn=chat_fn,
                additional_inputs=[session_state],
                chatbot=gr.Chatbot(height=520),
                textbox=gr.Textbox(
                    placeholder="Ask a question about the indexed documents or Tazim's CV...",
                    container=False,
                    scale=7,
                ),
            )
        with gr.TabItem("📁 Add Knowledge"):
            gr.Markdown(
                "### Drop your files below to index them\n"
                "Supported formats: **PDF, DOCX, TXT, MD, CSV, JSON** "
                "(and most plain-text formats).  \n"
                "The filename is automatically used as the source name."
            )
            file_drop = gr.File(
                label="Drop files here or click to upload",
                file_count="multiple",
                type="filepath",
            )
            ingest_btn    = gr.Button("Ingest Files", variant="primary")
            ingest_status = gr.Textbox(label="Ingestion Status", interactive=False, lines=4)

            gr.Markdown("---")
            gr.Markdown("### 📚 Previously Uploaded & Indexed Files")
            indexed_files_display = gr.Markdown(value=get_indexed_sources_markdown())
            refresh_files_btn = gr.Button("🔄 Refresh Indexed Files List", variant="secondary")


            with gr.Row():
                delete_input = gr.Textbox(label="Delete File / Source Name", placeholder="Enter exact source filename (e.g., Tazim_CV.pdf)...", scale=4)
                delete_btn = gr.Button("🗑️ Delete Source", variant="stop", scale=1)
                delete_status = gr.Textbox(label="Delete Status", interactive=False, scale=2)

            ingest_btn.click(fn=ingest_file_and_refresh, inputs=[file_drop], outputs=[ingest_status, indexed_files_display])
            refresh_files_btn.click(fn=get_indexed_sources_markdown, inputs=[], outputs=[indexed_files_display])
            delete_btn.click(fn=delete_indexed_source_fn, inputs=[delete_input], outputs=[delete_status, indexed_files_display])


        with gr.TabItem("⚡ System Dashboard & Architecture"):
            gr.Markdown("### Real-Time System Architecture, OTel Latency Waterfall & Database Inspector")
            gr.HTML(
                '<iframe src="/dashboard" style="width:100%; height:820px; border:1px solid rgba(255,255,255,0.1); border-radius:14px;"></iframe>'
            )

# Mount Gradio onto the FastAPI app (matches mini_rag.py)
app = gr.mount_gradio_app(app, demo, path="/")


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
